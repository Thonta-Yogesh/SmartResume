"""
SmartResume Engine — Local NLP-powered resume processing and semantic search.

Uses SentenceTransformers (all-mpnet-base-v2) for embedding generation
and cosine similarity for semantic matching. No external APIs required.

NLP Pipeline (as per research paper):
  - Tokenization, stopword removal, lemmatization via NLTK
  - Named Entity Recognition (NER) via spaCy
  - Semantic embeddings via Sentence Transformers
"""

import os
import re
import json
import hashlib
import logging
import numpy as np
from typing import List, Dict, Any, Optional
from datetime import datetime

# Text extraction
import PyPDF2
from docx import Document

# NLP Preprocessing (Paper Section III-B)
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# Named Entity Recognition (Paper Section IV-B)
import spacy

# NLP Embeddings
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# ─── Logging Setup ────────────────────────────────────────────
logger = logging.getLogger("SmartResume.Engine")
logger.setLevel(logging.INFO)
if not logger.handlers:
    _handler = logging.StreamHandler()
    _handler.setFormatter(logging.Formatter(
        "%(asctime)s [%(name)s] %(levelname)s: %(message)s",
        datefmt="%H:%M:%S"
    ))
    logger.addHandler(_handler)

# ─── NLTK / spaCy bootstrap ──────────────────────────────────
for _res in ["punkt_tab", "stopwords", "wordnet"]:
    try:
        nltk.data.find(f"tokenizers/{_res}" if "punkt" in _res
                       else f"corpora/{_res}")
    except LookupError:
        nltk.download(_res, quiet=True)

try:
    _nlp = spacy.load("en_core_web_sm")
except OSError:
    from spacy.cli import download as _spacy_dl
    _spacy_dl("en_core_web_sm")
    _nlp = spacy.load("en_core_web_sm")


# ─── Constants ────────────────────────────────────────────────
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STORE_PATH = os.path.join(_BASE_DIR, "resume_store.json")
UPLOADS_DIR = os.path.join(_BASE_DIR, "uploaded_resumes")
CHUNK_SIZE = 1500  # characters per chunk
EMBEDDING_DIM = 768

os.makedirs(UPLOADS_DIR, exist_ok=True)

# Common tech keywords for skill extraction
TECH_KEYWORDS = [
    # Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "perl",
    # Web
    "react", "angular", "vue", "node.js", "express", "django", "flask",
    "spring", "html", "css", "bootstrap", "tailwind", "next.js", "fastapi",
    # Data / ML
    "machine learning", "deep learning", "nlp", "natural language processing",
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
    "data science", "data analysis", "data engineering", "big data",
    "computer vision", "neural network", "ai", "artificial intelligence",
    # Cloud / DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "jenkins",
    "ci/cd", "terraform", "ansible", "linux", "devops", "microservices",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "dynamodb", "cassandra", "oracle", "firebase", "supabase",
    # Tools
    "git", "github", "jira", "agile", "scrum", "rest api", "graphql",
    "power bi", "tableau", "excel", "figma", "adobe",
    # Soft skills / roles
    "project management", "team lead", "leadership", "communication",
    "problem solving", "full stack", "frontend", "backend", "software engineer",
    "web developer", "data scientist", "data analyst", "ml engineer",
]


class ResumeEngine:
    """
    Self-contained resume processing and semantic search engine.
    Stores all data locally in a JSON file.

    NLP pipeline (aligned with research paper methodology):
      1. Text extraction  (PyPDF2 / python-docx)
      2. NLP preprocessing (tokenization → stopword removal → lemmatization)
      3. Text chunking (paragraph-aware dynamic chunking)
      4. Embedding generation (Sentence Transformers all-mpnet-base-v2, 768-d)
      5. Semantic search (cosine similarity + skill bonus)
    """

    def __init__(self, model_name: str = "all-mpnet-base-v2"):
        self.model_name = model_name
        self._model = None  # Lazy loading
        self._lemmatizer = WordNetLemmatizer()
        self._stop_words = set(stopwords.words("english"))
        self.store: Dict[str, Any] = {"resumes": {}}
        self._load_store()

    @property
    def model(self) -> SentenceTransformer:
        """Lazy-load the embedding model."""
        if self._model is None:
            self._model = SentenceTransformer(self.model_name)
        return self._model

    # ─── Store Management ─────────────────────────────────────

    def _load_store(self):
        """Load the resume store from disk."""
        if os.path.exists(STORE_PATH):
            try:
                with open(STORE_PATH, "r", encoding="utf-8") as f:
                    self.store = json.load(f)
            except (json.JSONDecodeError, IOError):
                self.store = {"resumes": {}}
        else:
            self.store = {"resumes": {}}

    def _save_store(self):
        """Persist the resume store to disk."""
        with open(STORE_PATH, "w", encoding="utf-8") as f:
            json.dump(self.store, f, indent=2, ensure_ascii=False)

    # ─── Text Extraction ──────────────────────────────────────

    @staticmethod
    def extract_text(file_path: str = None, file_bytes: bytes = None,
                     file_name: str = "") -> str:
        """Extract text from PDF, DOCX, or TXT files."""
        name = (file_name or file_path or "").lower()

        if name.endswith(".pdf"):
            return ResumeEngine._extract_pdf(file_path, file_bytes)
        elif name.endswith(".docx"):
            return ResumeEngine._extract_docx(file_path, file_bytes)
        elif name.endswith(".txt"):
            if file_bytes:
                return file_bytes.decode("utf-8", errors="ignore").strip()
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        else:
            raise ValueError(f"Unsupported file type: {name}")

    @staticmethod
    def _extract_pdf(file_path=None, file_bytes=None) -> str:
        import io
        text = ""
        try:
            if file_bytes:
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            else:
                reader = PyPDF2.PdfReader(file_path)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        except Exception as e:
            logger.error("Error reading PDF: %s", e)
        return text.strip()

    @staticmethod
    def _extract_docx(file_path=None, file_bytes=None) -> str:
        import io
        text = ""
        try:
            if file_bytes:
                doc = Document(io.BytesIO(file_bytes))
            else:
                doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error("Error reading DOCX: %s", e)
        return text.strip()

    # ─── NLP Preprocessing (Paper §III-B) ─────────────────────

    def preprocess_text(self, text: str) -> str:
        """
        Full NLP preprocessing pipeline as described in the research paper:
          1. Tokenization — split text into individual word tokens
          2. Stopword removal — remove non-informative words ("the", "is", "and")
          3. Lemmatization — reduce words to base forms ("running" → "run")

        Returns cleaned, lemmatized text ready for embedding generation.
        """
        # Step 1: Tokenize
        tokens = word_tokenize(text.lower())

        # Step 2: Remove stopwords and non-alphabetic tokens
        filtered = [
            tok for tok in tokens
            if tok.isalpha() and tok not in self._stop_words
        ]

        # Step 3: Lemmatize
        lemmatized = [self._lemmatizer.lemmatize(tok) for tok in filtered]

        return " ".join(lemmatized)

    # ─── Text Chunking ────────────────────────────────────────

    @staticmethod
    def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> List[str]:
        """Split text into overlapping chunks at paragraph boundaries."""
        chunks = []
        start = 0
        length = len(text)

        while start < length:
            end = min(start + chunk_size, length)
            if end < length:
                # Try to break at paragraph boundary
                chunk = text[start:end]
                last_break = chunk.rfind("\n\n")
                if last_break > chunk_size * 0.3:
                    end = start + last_break
            chunks.append(text[start:end].strip())
            start = end

        return [c for c in chunks if c]

    @staticmethod
    def extract_candidate_name(text: str) -> str:
        """Try to extract candidate name from resume text (first non-empty line)."""
        lines = text.strip().split("\n")
        for line in lines[:5]:
            clean = line.strip()
            # Name is usually the first short line without common headers
            if clean and len(clean) < 60 and not any(
                kw in clean.lower() for kw in
                ["resume", "curriculum", "objective", "summary", "email",
                 "phone", "address", "http", "www", "@", "linkedin"]
            ):
                return clean
        return "Unknown Candidate"

    @staticmethod
    def extract_skills(text: str) -> List[str]:
        """Extract technical skills from resume text using keyword matching."""
        text_lower = text.lower()
        found = []
        for kw in TECH_KEYWORDS:
            if kw in text_lower:
                found.append(kw)
        return sorted(set(found))

    @staticmethod
    def extract_skills_ner(text: str) -> List[str]:
        """
        Extract skills using spaCy Named Entity Recognition (NER).
        As described in the paper (Section IV-B): "Named Entity Recognition
        (NER): Classifying and categorizing resumes according to crucial
        competencies, technologies and skills."

        Extracts ORG, PRODUCT, and LANGUAGE entities that commonly
        correspond to technologies, frameworks, and programming languages.
        """
        doc = _nlp(text)
        ner_skills = set()
        # Relevant entity labels for tech skills
        tech_labels = {"ORG", "PRODUCT", "LANGUAGE", "WORK_OF_ART"}
        for ent in doc.ents:
            if ent.label_ in tech_labels and len(ent.text) > 1:
                ner_skills.add(ent.text.strip())
        return sorted(ner_skills)

    @staticmethod
    def extract_experience_years(text: str) -> float:
        """Extract approximate years of experience from resume text."""
        total = 0.0

        # Pattern: "X years" / "X+ years of experience"
        year_matches = re.findall(
            r"(\d+)\+?\s*(?:years?|yrs?)(?:\s+of\s+experience)?",
            text, re.IGNORECASE
        )
        if year_matches:
            total = max(float(y) for y in year_matches)

        # Pattern: date ranges "2019 - 2024" or "2019 – Present"
        date_ranges = re.findall(
            r"(20\d{2}|19\d{2})\s*[-–—]\s*(20\d{2}|19\d{2}|[Pp]resent|[Cc]urrent)",
            text
        )
        for start_str, end_str in date_ranges:
            start = int(start_str)
            end = datetime.now().year if end_str.lower() in ("present", "current") else int(end_str)
            span = end - start
            if 0 < span <= 40:
                total = max(total, float(span))

        return total

    @staticmethod
    def extract_email(text: str) -> str:
        """Extract email from resume."""
        match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
        return match.group(0) if match else ""

    @staticmethod
    def extract_phone(text: str) -> str:
        """Extract phone number from resume."""
        match = re.search(r"[\+]?[\d\s\-\(\)]{10,15}", text)
        return match.group(0).strip() if match else ""

    # ─── Embedding & Search ───────────────────────────────────

    def generate_embedding(self, text: str) -> List[float]:
        """Generate a 768-dim embedding vector for text."""
        embedding = self.model.encode(text, show_progress_bar=False)
        if isinstance(embedding, np.ndarray):
            return embedding.tolist()
        return list(embedding)

    def process_resume(self, file_name: str, file_bytes: bytes,
                       progress_callback=None) -> Dict[str, Any]:
        """
        Process a single resume file:
        1. Extract text
        2. Extract metadata (name, skills, experience, contact)
        3. Chunk text
        4. Generate embeddings for each chunk
        5. Store in local JSON database
        """
        # Generate unique ID
        file_hash = hashlib.md5(file_bytes).hexdigest()[:12]
        resume_id = f"{file_hash}_{file_name}"

        # Check if already processed
        if resume_id in self.store["resumes"]:
            return self.store["resumes"][resume_id]

        # Step 0: Save original file to disk for later download
        saved_path = os.path.join(UPLOADS_DIR, file_name)
        with open(saved_path, "wb") as f:
            f.write(file_bytes)

        # Step 1: Extract text
        if progress_callback:
            progress_callback(0.1, "Extracting text...")
        text = self.extract_text(file_bytes=file_bytes, file_name=file_name)
        if not text:
            raise ValueError(f"Could not extract text from {file_name}")

        # Step 1.5: NLP Preprocessing (Paper §III-B)
        # Tokenization → stopword removal → lemmatization
        if progress_callback:
            progress_callback(0.2, "NLP preprocessing (tokenize, stopwords, lemmatize)...")
        preprocessed_text = self.preprocess_text(text)
        logger.info("Preprocessed %s: %d chars → %d tokens",
                    file_name, len(text), len(preprocessed_text.split()))

        # Step 2: Extract metadata
        if progress_callback:
            progress_callback(0.3, "Analyzing resume content...")
        candidate_name = self.extract_candidate_name(text)
        skills_keyword = self.extract_skills(text)
        skills_ner = self.extract_skills_ner(text)
        # Merge both skill-extraction methods (keyword + NER)
        skills = sorted(set(skills_keyword) | {s.lower() for s in skills_ner})
        experience = self.extract_experience_years(text)
        email = self.extract_email(text)
        phone = self.extract_phone(text)
        logger.info("Extracted %d keyword skills + %d NER entities for %s",
                    len(skills_keyword), len(skills_ner), file_name)

        # Step 3: Chunk text
        if progress_callback:
            progress_callback(0.5, "Chunking text...")
        # Use preprocessed text for embedding (better vector quality)
        # but keep original text for display / human-readable output
        chunks = self.chunk_text(text)

        # Step 4: Generate embeddings on preprocessed text
        #   Use preprocessed (lemmatized) chunks for embedding generation
        #   so that the vectors capture normalized semantic meaning.
        preprocessed_chunks = [self.preprocess_text(c) for c in chunks]
        if progress_callback:
            progress_callback(0.6, "Generating embeddings (NLP)...")
        chunk_embeddings = []
        for i, pchunk in enumerate(preprocessed_chunks):
            emb = self.generate_embedding(pchunk)
            chunk_embeddings.append(emb)
            if progress_callback:
                prog = 0.6 + 0.3 * ((i + 1) / len(chunks))
                progress_callback(prog, f"Embedding chunk {i+1}/{len(chunks)}...")

        # Step 5: Generate a full-document embedding (for quick overview matching)
        #   Use preprocessed text for the full-document embedding as well
        full_embedding = self.generate_embedding(preprocessed_text[:5000])

        # Build resume record
        resume_data = {
            "id": resume_id,
            "file_name": file_name,
            "file_path": saved_path,
            "candidate_name": candidate_name,
            "skills": skills,
            "experience_years": experience,
            "email": email,
            "phone": phone,
            "full_text": text,
            "chunks": chunks,
            "chunk_embeddings": chunk_embeddings,
            "full_embedding": full_embedding,
            "skills_ner": skills_ner,
            "processed_at": datetime.now().isoformat(),
        }

        # Save to store
        self.store["resumes"][resume_id] = resume_data
        self._save_store()

        if progress_callback:
            progress_callback(1.0, "Done!")

        return resume_data

    def search(self, query: str, top_k: int = 10) -> List[Dict[str, Any]]:
        """
        Semantic search: find the best matching resumes for a job description/query.

        Uses cosine similarity between query embedding and each resume's
        full-document embedding + best chunk embedding.
        """
        if not self.store["resumes"]:
            return []

        # Generate query embedding
        query_embedding = np.array(self.generate_embedding(query)).reshape(1, -1)

        results = []
        for rid, resume in self.store["resumes"].items():
            # Full-document similarity
            full_emb = np.array(resume["full_embedding"]).reshape(1, -1)
            full_sim = cosine_similarity(query_embedding, full_emb)[0][0]

            # Best chunk similarity (for more granular matching)
            best_chunk_sim = 0.0
            best_chunk_idx = 0
            for i, cemb in enumerate(resume["chunk_embeddings"]):
                chunk_emb = np.array(cemb).reshape(1, -1)
                sim = cosine_similarity(query_embedding, chunk_emb)[0][0]
                if sim > best_chunk_sim:
                    best_chunk_sim = sim
                    best_chunk_idx = i

            # Combined score: 60% full-doc + 40% best chunk
            combined_score = 0.6 * full_sim + 0.4 * best_chunk_sim

            # Skill match bonus
            query_lower = query.lower()
            query_skills = [kw for kw in TECH_KEYWORDS if kw in query_lower]
            resume_skills = resume.get("skills", [])
            if query_skills:
                skill_overlap = len(set(query_skills) & set(resume_skills))
                skill_bonus = 0.1 * (skill_overlap / len(query_skills))
                combined_score += skill_bonus

            results.append({
                "resume_id": rid,
                "candidate_name": resume["candidate_name"],
                "file_name": resume["file_name"],
                "file_path": resume.get("file_path", ""),
                "match_score": round(float(combined_score) * 100, 1),
                "full_doc_score": round(float(full_sim) * 100, 1),
                "chunk_score": round(float(best_chunk_sim) * 100, 1),
                "best_chunk": resume["chunks"][best_chunk_idx] if resume["chunks"] else "",
                "skills": resume_skills,
                "experience_years": resume.get("experience_years", 0),
                "email": resume.get("email", ""),
                "phone": resume.get("phone", ""),
            })

        # Sort by combined score descending
        results.sort(key=lambda x: x["match_score"], reverse=True)
        return results[:top_k]

    def get_all_resumes(self) -> List[Dict[str, Any]]:
        """Return summary of all stored resumes."""
        summaries = []
        for rid, resume in self.store["resumes"].items():
            summaries.append({
                "id": rid,
                "file_name": resume["file_name"],
                "candidate_name": resume["candidate_name"],
                "skills": resume.get("skills", []),
                "experience_years": resume.get("experience_years", 0),
                "email": resume.get("email", ""),
                "processed_at": resume.get("processed_at", ""),
            })
        return summaries

    def delete_resume(self, resume_id: str):
        """Delete a resume from the store."""
        if resume_id in self.store["resumes"]:
            del self.store["resumes"][resume_id]
            self._save_store()

    def clear_all(self):
        """Clear all stored resumes."""
        self.store = {"resumes": {}}
        self._save_store()

    def get_resume_file_bytes(self, resume_id: str) -> tuple:
        """Return (file_name, file_bytes) for a resume, or None if not found."""
        resume = self.store["resumes"].get(resume_id)
        if not resume:
            return None, None
        file_path = resume.get("file_path", "")
        # Try stored path first, then check local resumes folder
        if file_path and os.path.exists(file_path):
            with open(file_path, "rb") as f:
                return resume["file_name"], f.read()
        # Fallback: check resumes/ folder
        alt_path = os.path.join(_BASE_DIR, "resumes", resume["file_name"])
        if os.path.exists(alt_path):
            with open(alt_path, "rb") as f:
                return resume["file_name"], f.read()
        return resume["file_name"], None

    @property
    def resume_count(self) -> int:
        return len(self.store["resumes"])
